from pathlib import Path
from docx import Document

root = Path(r'd:\image-enhancement-project\AI-Powered-Image-Super-Resolution-and-Detail-Enhancement-System')
md_path = root / 'AWS_DEPLOYMENT_GUIDE.md'
out_path = root / 'AWS_DEPLOYMENT_GUIDE.docx'

text = md_path.read_text(encoding='utf-8')
lines = text.splitlines()

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = 11

in_code = False
code_buffer = []

for line in lines:
    if line.startswith('```'):
        if in_code:
            p = doc.add_paragraph()
            run = p.add_run('\n'.join(code_buffer))
            run.font.name = 'Consolas'
            run.font.size = 10
            code_buffer = []
            in_code = False
        else:
            in_code = True
        continue

    if in_code:
        code_buffer.append(line)
        continue

    if line.startswith('# '):
        doc.add_heading(line[2:], level=1)
    elif line.startswith('## '):
        doc.add_heading(line[3:], level=2)
    elif line.startswith('### '):
        doc.add_heading(line[4:], level=3)
    elif line.startswith('- '):
        doc.add_paragraph(line[2:], style='List Bullet')
    elif line.startswith(('1. ', '2. ', '3. ', '4. ', '5. ')):
        doc.add_paragraph(line, style='List Number')
    elif line.strip():
        doc.add_paragraph(line)

if in_code:
    p = doc.add_paragraph()
    run = p.add_run('\n'.join(code_buffer))
    run.font.name = 'Consolas'
    run.font.size = 10

doc.save(out_path)
print(out_path)
